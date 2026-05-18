"""
Word文档解析模块 - 读取论文的格式信息
"""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, asdict
import re


@dataclass
class ParagraphInfo:
    """段落信息"""
    index: int
    text: str
    style_name: str
    font_name: str
    font_size: float
    bold: bool
    italic: bool
    alignment: str
    line_spacing: float
    first_line_indent: float  # 首行缩进（字符数）
    space_before: float
    space_after: float
    section_type: str  # title/abstract/body/reference/figure/table/unknown


@dataclass
class TableInfo:
    """表格信息"""
    index: int
    rows: int
    cols: int
    caption: str
    is_float: bool


@dataclass
class ImageInfo:
    """图片信息"""
    index: int
    width: float
    height: float
    caption: str
    position: str


@dataclass
class DocumentStructure:
    """文档结构"""
    paragraphs: List[ParagraphInfo]
    tables: List[TableInfo]
    images: List[ImageInfo]
    section_mapping: Dict[str, List[int]]  # section_type -> paragraph_indices


class DocxParser:
    """Word文档解析器"""
    
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.structure: Optional[DocumentStructure] = None
        
    def parse(self) -> DocumentStructure:
        """解析整个文档"""
        paragraphs = self._parse_paragraphs()
        tables = self._parse_tables()
        images = self._parse_images()
        section_mapping = self._map_sections(paragraphs)
        
        self.structure = DocumentStructure(
            paragraphs=paragraphs,
            tables=tables,
            images=images,
            section_mapping=section_mapping
        )
        return self.structure
    
    def _parse_paragraphs(self) -> List[ParagraphInfo]:
        """解析所有段落"""
        paragraphs = []
        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:  # 跳过空段落
                continue
                
            # 获取格式信息
            font_name = "宋体"
            font_size = 12.0
            bold = False
            italic = False
            
            if para.runs:
                run = para.runs[0]
                # 字体名称
                if run.font.name:
                    font_name = run.font.name
                # 字号
                if run.font.size:
                    font_size = run.font.size.pt
                # 加粗/斜体
                bold = run.bold or (para.style.font.bold if para.style.font.bold is not None else False)
                italic = run.italic or (para.style.font.italic if para.style.font.italic is not None else False)
            
            # 对齐方式
            alignment_map = {
                WD_ALIGN_PARAGRAPH.LEFT: "left",
                WD_ALIGN_PARAGRAPH.CENTER: "center",
                WD_ALIGN_PARAGRAPH.RIGHT: "right",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
            }
            alignment = alignment_map.get(para.alignment, "left")
            
            # 行距
            line_spacing = 1.5  # 默认1.5倍行距
            if para.paragraph_format.line_spacing:
                line_spacing = para.paragraph_format.line_spacing
            
            # 首行缩进（转换为字符数，假设1字符约10.5pt）
            first_line_indent = 0
            if para.paragraph_format.first_line_indent:
                first_line_indent = para.paragraph_format.first_line_indent.pt / 10.5
            
            # 段前段后间距
            space_before = 0
            space_after = 0
            if para.paragraph_format.space_before:
                space_before = para.paragraph_format.space_before.pt
            if para.paragraph_format.space_after:
                space_after = para.paragraph_format.space_after.pt
            
            # 识别段落类型
            section_type = self._detect_section_type(text, para.style_name)
            
            paragraphs.append(ParagraphInfo(
                index=idx,
                text=text,
                style_name=para.style_name or "Normal",
                font_name=font_name,
                font_size=font_size,
                bold=bold,
                italic=italic,
                alignment=alignment,
                line_spacing=line_spacing,
                first_line_indent=first_line_indent,
                space_before=space_before,
                space_after=space_after,
                section_type=section_type
            ))
        
        return paragraphs
    
    def _detect_section_type(self, text: str, style_name: str) -> str:
        """识别段落类型"""
        text_lower = text.lower()
        style_lower = style_name.lower() if style_name else ""
        
        # 基于样式名称判断
        if "title" in style_lower or "标题" in style_lower:
            return "title"
        if "heading" in style_lower or "正文" in style_lower or "body" in style_lower:
            return "body"
        if "abstract" in style_lower or "摘要" in style_lower:
            return "abstract"
        if "reference" in style_lower or "reference" in style_lower:
            return "reference"
        
        # 基于关键词判断
        if re.match(r'^\d+[\.、]\s*\S+', text) and len(text) < 50:  # 标题样式编号
            return "body"
        if re.match(r'^\[?\d+\]?\s+[\u4e00-\u9fa5]', text) or re.match(r'^\[\d+\]', text):
            # 参考文献条目
            if any(kw in text_lower for kw in ['doi', 'http', '出版社', 'journal', 'conference']):
                return "reference"
        
        return "unknown"
    
    def _parse_tables(self) -> List[TableInfo]:
        """解析表格"""
        tables = []
        for idx, table in enumerate(self.doc.tables):
            rows = len(table.rows)
            cols = len(table.columns) if table.rows else 0
            
            # 尝试获取标题
            caption = ""
            # 检查前后段落是否有"表X"或"Table X"
            if self.structure and idx > 0:
                for p in reversed(self.structure.paragraphs[:idx]):
                    if re.search(r'表\s*\d|Table\s*\d', p.text):
                        caption = p.text
                        break
            
            tables.append(TableInfo(
                index=idx,
                rows=rows,
                cols=cols,
                caption=caption,
                is_float=True
            ))
        return tables
    
    def _parse_images(self) -> List[ImageInfo]:
        """解析图片（简化版，实际需要遍历inline shapes）"""
        images = []
        # 遍历文档中的所有内联形状
        for para in self.doc.paragraphs:
            for run in para.runs:
                # 检查run中是否有图片
                pass  # 简化处理
        return images
    
    def _map_sections(self, paragraphs: List[ParagraphInfo]) -> Dict[str, List[int]]:
        """建立段落与章节的映射"""
        mapping: Dict[str, List[int]] = {
            "title": [],
            "abstract": [],
            "keywords": [],
            "body": [],
            "reference": [],
            "acknowledgment": [],
            "appendix": [],
            "unknown": []
        }
        
        current_section = "unknown"
        
        for idx, para in enumerate(paragraphs):
            text_lower = para.text.lower()
            
            # 检测章节变化
            if re.match(r'^(摘要|abstract)', text_lower):
                current_section = "abstract"
            elif re.match(r'^(关键词|关键字|keywords)', text_lower):
                current_section = "keywords"
            elif re.match(r'^(引言|前言|背景|1\s)', text_lower):
                current_section = "body"
            elif re.match(r'^(参考文献|reference)', text_lower):
                current_section = "reference"
            elif re.match(r'^(致谢|acknowledg)', text_lower):
                current_section = "acknowledgment"
            elif re.match(r'^(附录|appendix)', text_lower):
                current_section = "appendix"
            
            mapping[current_section].append(idx)
        
        return mapping
    
    def get_text_content(self) -> str:
        """获取纯文本内容"""
        return "\n".join([p.text for p in self.structure.paragraphs])
    
    def get_format_summary(self) -> Dict[str, Any]:
        """获取格式摘要"""
        if not self.structure:
            self.parse()
        
        summary = {
            "total_paragraphs": len(self.structure.paragraphs),
            "total_tables": len(self.structure.tables),
            "total_images": len(self.structure.images),
            "sections": {}
        }
        
        for section_type, indices in self.structure.section_mapping.items():
            if indices:
                summary["sections"][section_type] = len(indices)
        
        return summary


def extract_docx_format(docx_path: str) -> Dict[str, Any]:
    """便捷函数：提取文档格式"""
    parser = DocxParser(docx_path)
    structure = parser.parse()
    
    return {
        "structure": structure,
        "summary": parser.get_format_summary(),
        "full_text": parser.get_text_content()
    }
