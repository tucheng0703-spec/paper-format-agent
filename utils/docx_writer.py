"""
Word文档生成模块 - 应用格式规则生成新文档
"""
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Any, Optional
import copy

from utils.format_engine import FormatRule, RuleSet
from utils.docx_parser import DocumentStructure, ParagraphInfo


class DocxWriter:
    """Word文档生成器"""
    
    def __init__(self):
        self.doc = Document()
        self.rules: Optional[RuleSet] = None
        
    def load_rules(self, rules: RuleSet):
        """加载规则集"""
        self.rules = rules
    
    def load_rules_from_json(self, json_str: str):
        """从JSON加载规则"""
        self.rules = RuleSet.from_json(json_str)
    
    def apply_format_to_run(self, run, rule: FormatRule):
        """应用规则到run（文字片段）"""
        if rule.font_name:
            run.font.name = rule.font_name
            # 设置中文字体
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), rule.font_name)
        
        if rule.font_size is not None:
            run.font.size = Pt(rule.font_size)
        
        if rule.bold is not None:
            run.font.bold = rule.bold
        
        if rule.italic is not None:
            run.font.italic = rule.italic
    
    def apply_format_to_paragraph(self, para, rule: FormatRule):
        """应用规则到paragraph"""
        pf = para.paragraph_format
        
        # 对齐方式
        if rule.alignment:
            alignment_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            pf.alignment = alignment_map.get(rule.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        # 行距
        if rule.line_spacing is not None:
            pf.line_spacing = rule.line_spacing
        
        # 首行缩进（字符数转Pt，1字符约10.5pt）
        if rule.first_line_indent is not None:
            pf.first_line_indent = Pt(rule.first_line_indent * 10.5)
        
        # 段前段后间距
        if rule.space_before is not None:
            pf.space_before = Pt(rule.space_before)
        
        if rule.space_after is not None:
            pf.space_after = Pt(rule.space_after)
    
    def get_rules_for_section(self, section_type: str) -> FormatRule:
        """获取某章节类型的规则（合并所有适用规则）"""
        if not self.rules:
            return FormatRule(target=section_type)
        
        rules = self.rules.get_rules_for_target(section_type)
        if not rules:
            # 尝试使用正文规则
            rules = self.rules.get_rules_for_target("body")
        
        if not rules:
            return FormatRule(target=section_type)
        
        # 合并所有规则
        merged = FormatRule(target=section_type)
        for rule in rules:
            for key in ["font_name", "font_size", "bold", "italic", "alignment", 
                       "line_spacing", "first_line_indent", "space_before", "space_after"]:
                value = getattr(rule, key)
                if value is not None:
                    setattr(merged, key, value)
        
        return merged
    
    def add_paragraph_with_format(self, text: str, section_type: str, preserve_style: bool = False):
        """添加带格式的段落"""
        para = self.doc.add_paragraph()
        
        # 获取适用的规则
        rule = self.get_rules_for_section(section_type)
        
        # 添加文本
        run = para.add_run(text)
        
        # 应用规则
        self.apply_format_to_run(run, rule)
        self.apply_format_to_paragraph(para, rule)
        
        return para
    
    def create_from_structure(self, structure: DocumentStructure) -> Document:
        """从文档结构创建新文档"""
        for para_info in structure.paragraphs:
            self.add_paragraph_with_format(
                para_info.text, 
                para_info.section_type
            )
        
        return self.doc
    
    def save(self, output_path: str):
        """保存文档"""
        self.doc.save(output_path)


class FormatApplier:
    """格式应用器 - 将规则应用到已有文档"""
    
    def __init__(self):
        self.rules: Optional[RuleSet] = None
        self.changes: List[Dict] = []  # 记录修改
    
    def load_rules(self, rules: RuleSet):
        self.rules = rules
    
    def apply_to_document(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """将规则应用到文档"""
        from utils.docx_parser import DocxParser
        
        # 解析原文档
        parser = DocxParser(input_path)
        structure = parser.parse()
        
        # 创建新文档
        writer = DocxWriter()
        writer.rules = self.rules
        
        # 逐段处理
        for para_info in structure.paragraphs:
            para = writer.doc.add_paragraph()
            
            # 添加文本
            run = para.add_run(para_info.text)
            
            # 获取该部分的规则
            rule = writer.get_rules_for_section(para_info.section_type)
            
            # 应用格式
            writer.apply_format_to_run(run, rule)
            writer.apply_format_to_paragraph(para, rule)
            
            # 记录修改
            self.changes.append({
                "index": para_info.index,
                "text": para_info.text[:50],
                "section": para_info.section_type,
                "new_format": rule.to_dict()
            })
        
        # 保存
        writer.save(output_path)
        
        return {
            "total_paragraphs": len(structure.paragraphs),
            "applied_rules": self.rules.to_json() if self.rules else "{}",
            "changes": self.changes
        }


def apply_rules_to_docx(input_path: str, output_path: str, rules: RuleSet) -> Dict[str, Any]:
    """便捷函数：应用规则到docx文件"""
    applier = FormatApplier()
    applier.load_rules(rules)
    return applier.apply_to_document(input_path, output_path)
